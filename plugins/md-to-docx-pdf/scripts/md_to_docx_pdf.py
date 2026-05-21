#!/usr/bin/env python3
"""Render Markdown to a styled DOCX and PDF (PDF via LibreOffice headless).

A general-purpose document builder: parses a practical subset of Markdown
(headings, inline bold/italic/code with escapes, tables with content-weighted
auto-width columns, bullet/numbered/checkbox lists, fenced code blocks,
horizontal rules, embedded images) into a python-docx document with a title
page, then converts to PDF with LibreOffice.

Two ways to drive it:

  1. Single document from flags:
       md_to_docx_pdf.py INPUT.md [INPUT2.md ...] \
         --out-docx OUT.docx [--out-pdf OUT.pdf] \
         --title "Title" [--subtitle "..."] [--author "..."] \
         [--institute "..."] [--date YYYY-MM-DD] [--banner CONFIDENTIAL] \
         [--footer "line one" --footer "line two"]

  2. Multiple documents from a JSON config (build several in one run):
       md_to_docx_pdf.py --config build.json
     where build.json is {"documents": [ {<one build spec>}, ... ]}; each spec
     uses the same keys as the flags (inputs, out_docx, out_pdf, title, ...).

Multiple INPUT.md files concatenate into one document (page break between).
Images in `![alt](path)` resolve relative to each Markdown file's directory.

Requires: python-docx (`pip install python-docx`) and a `libreoffice` (or
`soffice`) binary on PATH for the PDF step. Set $LIBREOFFICE_BIN to override.
"""
from __future__ import annotations

import argparse
import datetime as _dt
import json
import os
import re
import shutil
import subprocess
import sys
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ==========================================================================
# Config
# ==========================================================================
@dataclass
class DocConfig:
    inputs: list[Path]
    out_docx: Path
    out_pdf: Path | None = None
    title: str = "Document"
    subtitle: str = ""
    author: str = ""
    institute: str = ""
    date: str = field(default_factory=lambda: _dt.date.today().isoformat())
    banner: str = ""               # e.g. "CONFIDENTIAL"; blank = no banner
    footer_lines: list[str] = field(default_factory=list)  # explicit overrides
    # styling
    body_font: str = "Calibri"
    body_pt: float = 10.5
    heading_color: str = "1A1A2E"  # hex
    page_w_cm: float = 21.0        # A4
    page_h_cm: float = 29.7
    margin_top_cm: float = 2.0
    margin_bottom_cm: float = 2.0
    margin_lr_cm: float = 2.5
    table_width_cm: float = 16.0
    image_width_cm: float = 15.5

    @property
    def resolved_pdf(self) -> Path:
        return self.out_pdf or self.out_docx.with_suffix(".pdf")


def _hex(rgb: str) -> RGBColor:
    rgb = rgb.lstrip("#")
    return RGBColor(int(rgb[0:2], 16), int(rgb[2:4], 16), int(rgb[4:6], 16))


# ==========================================================================
# Page / styles / title page
# ==========================================================================
def setup_styles(doc: Document, cfg: DocConfig) -> None:
    for section in doc.sections:
        section.page_width = Cm(cfg.page_w_cm)
        section.page_height = Cm(cfg.page_h_cm)
        section.top_margin = Cm(cfg.margin_top_cm)
        section.bottom_margin = Cm(cfg.margin_bottom_cm)
        section.left_margin = Cm(cfg.margin_lr_cm)
        section.right_margin = Cm(cfg.margin_lr_cm)

    style = doc.styles["Normal"]
    style.font.name = cfg.body_font
    style.font.size = Pt(cfg.body_pt)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)

    sizes = {1: 18, 2: 14, 3: 12, 4: 11}
    before = {1: 24, 2: 18, 3: 12, 4: 6}
    after = {1: 12, 2: 8, 3: 6, 4: 4}
    for level in range(1, 5):
        sname = f"Heading {level}"
        if sname in doc.styles:
            hs = doc.styles[sname]
            hs.font.name = cfg.body_font
            hs.font.color.rgb = _hex(cfg.heading_color)
            hs.font.size = Pt(sizes[level])
            hs.paragraph_format.space_before = Pt(before[level])
            hs.paragraph_format.space_after = Pt(after[level])


def add_title_page(doc: Document, cfg: DocConfig) -> None:
    for _ in range(5):
        doc.add_paragraph()

    if cfg.banner:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cfg.banner)
        run.font.size = Pt(14)
        run.bold = True
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(cfg.title)
    run.font.size = Pt(28)
    run.font.color.rgb = _hex(cfg.heading_color)
    run.bold = True

    if cfg.subtitle:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cfg.subtitle)
        run.font.size = Pt(16)
        run.font.color.rgb = _hex("555555")

    doc.add_paragraph()

    if cfg.footer_lines:
        footer_lines = list(cfg.footer_lines)
    else:
        footer_lines = []
        if cfg.author:
            footer_lines.append(f"Author: {cfg.author}")
        if cfg.institute:
            footer_lines.append(f"Institute: {cfg.institute}")
        footer_lines.append(f"Date: {cfg.date}")

    for line in footer_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = _hex("555555")

    doc.add_page_break()


# ==========================================================================
# Inline formatting (with markdown escapes \* and \\)
# ==========================================================================
_STAR_ESC = "\x01"
_BSLASH_ESC = "\x02"


def _unescape(s: str) -> str:
    return s.replace(_STAR_ESC, "*").replace(_BSLASH_ESC, "\\")


def add_inline_formatting(paragraph, text: str) -> None:
    """Parse inline markdown and append runs. Honors \\* and \\\\ escapes so
    literal characters survive the bold/italic/code regex."""
    text = text.replace("\\\\", _BSLASH_ESC).replace("\\*", _STAR_ESC)

    pattern = re.compile(
        r"(\*\*\*(.+?)\*\*\*)"   # bold-italic
        r"|(\*\*(.+?)\*\*)"      # bold
        r"|(\*(.+?)\*)"          # italic
        r"|(`([^`]+)`)"          # inline code
    )

    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(_unescape(text[pos:m.start()]))
        if m.group(2):
            run = paragraph.add_run(_unescape(m.group(2)))
            run.bold = True
            run.italic = True
        elif m.group(4):
            run = paragraph.add_run(_unescape(m.group(4)))
            run.bold = True
        elif m.group(6):
            run = paragraph.add_run(_unescape(m.group(6)))
            run.italic = True
        elif m.group(8):
            run = paragraph.add_run(_unescape(m.group(8)))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = _hex("880000")
        pos = m.end()

    if pos < len(text):
        paragraph.add_run(_unescape(text[pos:]))


# ==========================================================================
# Tables (content-weighted fixed column widths to minimise height)
# ==========================================================================
def _apply_table_layout(table, headers, rows, total_width_cm=16.0,
                        absolute_min_col_cm=0.8,
                        header_char_cm=0.25, content_char_cm=0.18):
    num_cols = len(headers)
    if num_cols == 0:
        return

    def _strip_md(s):
        return re.sub(r"\*\*|\*|`", "", s)

    def _max_atomic_word(s):
        atoms = re.split(r"[\s/\-]+", _strip_md(s))
        return max((len(a) for a in atoms), default=0)

    content_chars = [len(_strip_md(h)) for h in headers]
    cell_max_atom = [_max_atomic_word(h) for h in headers]
    for row_text in rows:
        cells_text = [c.strip() for c in row_text.strip().strip("|").split("|")]
        for i in range(min(num_cols, len(cells_text))):
            content_chars[i] += len(_strip_md(cells_text[i]))
            cell_max_atom[i] = max(cell_max_atom[i], _max_atomic_word(cells_text[i]))

    floors = [
        max(absolute_min_col_cm,
            _max_atomic_word(headers[i]) * header_char_cm,
            cell_max_atom[i] * content_char_cm)
        for i in range(num_cols)
    ]
    floor_total = sum(floors)

    if floor_total >= total_width_cm:
        scale = total_width_cm / floor_total
        widths_cm = [f * scale for f in floors]
    else:
        available = total_width_cm - floor_total
        total_weight = sum(content_chars) or num_cols
        widths_cm = [floors[i] + available * (content_chars[i] / total_weight)
                     for i in range(num_cols)]

    tbl = table._element
    tblPr = tbl.tblPr

    existing_layout = tblPr.find(qn("w:tblLayout"))
    if existing_layout is not None:
        tblPr.remove(existing_layout)
    tblPr.append(tblPr.makeelement(qn("w:tblLayout"), {qn("w:type"): "fixed"}))

    existing_w = tblPr.find(qn("w:tblW"))
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblPr.append(tblPr.makeelement(qn("w:tblW"), {
        qn("w:w"): str(int(Cm(total_width_cm).twips)),
        qn("w:type"): "dxa",
    }))

    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        for gc in list(tblGrid.findall(qn("w:gridCol"))):
            tblGrid.remove(gc)
        for w_cm in widths_cm:
            tblGrid.append(tblGrid.makeelement(qn("w:gridCol"),
                                               {qn("w:w"): str(int(Cm(w_cm).twips))}))

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < num_cols:
                cell.width = Cm(widths_cm[i])


def add_table(doc, header_line, rows, cfg: DocConfig):
    headers = [c.strip() for c in header_line.strip().strip("|").split("|")]
    num_cols = len(headers)

    table = doc.add_table(rows=1, cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = cfg.body_font
        shading = cell._element.get_or_add_tcPr()
        shading.append(shading.makeelement(qn("w:shd"),
                                           {qn("w:fill"): "E8E8F0", qn("w:val"): "clear"}))

    for row_text in rows:
        cells_text = [c.strip() for c in row_text.strip().strip("|").split("|")]
        row = table.add_row()
        for i in range(min(num_cols, len(cells_text))):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_formatting(p, cells_text[i])
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = cfg.body_font

    _apply_table_layout(table, headers, rows, total_width_cm=cfg.table_width_cm)
    doc.add_paragraph()


# ==========================================================================
# Markdown processing
# ==========================================================================
def process_markdown(doc, md_text, base_dir: Path, cfg: DocConfig, is_first=False):
    lines = md_text.split("\n")
    i = 0
    in_code_block = False
    h1_count = 0

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = _hex("333333")
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if re.match(r"^-{3,}$", line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("─" * 80)
            run.font.size = Pt(6)
            run.font.color.rgb = _hex("BBBBBB")
            i += 1
            continue

        heading_match = re.match(r"^(#{1,5})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            if level == 1:
                h1_count += 1
                if h1_count <= 1:   # first H1 handled by the title page
                    i += 1
                    continue
            p = doc.add_heading(level=min(level, 4))
            add_inline_formatting(p, text)
            i += 1
            continue

        if "|" in line and i + 1 < len(lines) and re.match(r"^\|?[\s\-:|]+\|", lines[i + 1]):
            header = line
            i += 2
            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                rows.append(lines[i])
                i += 1
            add_table(doc, header, rows, cfg)
            continue

        checkbox_match = re.match(r"^(\s*)- \[([ x])\]\s+(.+)$", line)
        if checkbox_match:
            indent = len(checkbox_match.group(1))
            checked = checkbox_match.group(2) == "x"
            marker = "☑" if checked else "☐"
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0 + (indent // 2) * 0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(f"{marker} ")
            run.font.size = Pt(10)
            add_inline_formatting(p, checkbox_match.group(3))
            i += 1
            continue

        bullet_match = re.match(r"^(\s*)[-*]\s+(.+)$", line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0 + (indent // 2) * 0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run("• ")
            run.font.size = Pt(10)
            add_inline_formatting(p, bullet_match.group(2))
            i += 1
            continue

        numbered_match = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if numbered_match:
            indent = len(numbered_match.group(1))
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0 + (indent // 3) * 0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            num = re.match(r"(\s*)(\d+)\.\s+", line).group(2)
            run = p.add_run(f"{num}. ")
            run.font.size = Pt(cfg.body_pt)
            add_inline_formatting(p, numbered_match.group(2))
            i += 1
            continue

        image_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip())
        if image_match:
            img_path = (base_dir / image_match.group(2)).resolve()
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.left_indent = Cm(0)
                p.paragraph_format.right_indent = Cm(0)
                p.add_run().add_picture(str(img_path), width=Cm(cfg.image_width_cm))
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"[Missing image: {img_path}]")
                run.italic = True
                run.font.color.rgb = _hex("CC0000")
            i += 1
            continue

        if line.strip().startswith("[Figure") or line.strip().startswith("[To be"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(line.strip())
            run.italic = True
            run.font.color.rgb = _hex("888888")
            run.font.size = Pt(10)
            i += 1
            continue

        if line.startswith("**") and ":" in line:
            p = doc.add_paragraph()
            add_inline_formatting(p, line)
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_formatting(p, line)
        i += 1

    return doc


# ==========================================================================
# Build
# ==========================================================================
def _libreoffice_bin() -> str | None:
    explicit = os.environ.get("LIBREOFFICE_BIN")
    if explicit and shutil.which(explicit):
        return explicit
    for cand in ("libreoffice", "soffice"):
        if shutil.which(cand):
            return cand
    return None


def build_document(cfg: DocConfig) -> None:
    doc = Document()
    setup_styles(doc, cfg)
    add_title_page(doc, cfg)

    for idx, filepath in enumerate(cfg.inputs):
        if not filepath.exists():
            print(f"Warning: {filepath} not found, skipping", file=sys.stderr)
            continue
        md_text = filepath.read_text(encoding="utf-8")
        if idx > 0:
            doc.add_page_break()
        process_markdown(doc, md_text, filepath.resolve().parent, cfg,
                         is_first=(idx == 0))

    cfg.out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(cfg.out_docx))
    print(f"DOCX saved: {cfg.out_docx}")

    lo = _libreoffice_bin()
    if not lo:
        print("WARNING: no libreoffice/soffice on PATH; skipped PDF. "
              "Set $LIBREOFFICE_BIN or install LibreOffice.", file=sys.stderr)
        return

    result = subprocess.run(
        [lo, "--headless", "--convert-to", "pdf",
         "--outdir", str(cfg.resolved_pdf.parent), str(cfg.out_docx)],
        capture_output=True, text=True,
    )
    if result.returncode != 0:
        print(f"PDF conversion failed: {result.stderr}", file=sys.stderr)
        sys.exit(1)
    # LibreOffice names the PDF after the DOCX stem; rename if a custom name was asked.
    produced = cfg.resolved_pdf.parent / (cfg.out_docx.stem + ".pdf")
    if produced != cfg.resolved_pdf and produced.exists():
        produced.replace(cfg.resolved_pdf)
    print(f"PDF saved: {cfg.resolved_pdf}")


# ==========================================================================
# CLI
# ==========================================================================
def _cfg_from_dict(d: dict, base: Path) -> DocConfig:
    def p(x):
        return (base / x) if not os.path.isabs(x) else Path(x)

    kwargs = dict(d)
    kwargs["inputs"] = [p(x) for x in d["inputs"]]
    kwargs["out_docx"] = p(d["out_docx"])
    if d.get("out_pdf"):
        kwargs["out_pdf"] = p(d["out_pdf"])
    return DocConfig(**kwargs)


def main(argv=None) -> None:
    ap = argparse.ArgumentParser(description="Render Markdown to styled DOCX + PDF.")
    ap.add_argument("inputs", nargs="*", help="Markdown input file(s); concatenated.")
    ap.add_argument("--config", help="JSON config building one or more documents.")
    ap.add_argument("--out-docx")
    ap.add_argument("--out-pdf")
    ap.add_argument("--title", default="Document")
    ap.add_argument("--subtitle", default="")
    ap.add_argument("--author", default="")
    ap.add_argument("--institute", default="")
    ap.add_argument("--date", default=_dt.date.today().isoformat())
    ap.add_argument("--banner", default="", help='e.g. "CONFIDENTIAL"')
    ap.add_argument("--footer", action="append", default=[],
                    help="Title-page footer line (repeatable; overrides auto footer).")
    args = ap.parse_args(argv)

    if args.config:
        cfg_path = Path(args.config).resolve()
        spec = json.loads(cfg_path.read_text())
        for doc_spec in spec["documents"]:
            build_document(_cfg_from_dict(doc_spec, cfg_path.parent))
        return

    if not args.inputs or not args.out_docx:
        ap.error("provide INPUT.md ... and --out-docx, or use --config")

    cfg = DocConfig(
        inputs=[Path(x).resolve() for x in args.inputs],
        out_docx=Path(args.out_docx).resolve(),
        out_pdf=Path(args.out_pdf).resolve() if args.out_pdf else None,
        title=args.title, subtitle=args.subtitle, author=args.author,
        institute=args.institute, date=args.date, banner=args.banner,
        footer_lines=args.footer,
    )
    build_document(cfg)


if __name__ == "__main__":
    main()
